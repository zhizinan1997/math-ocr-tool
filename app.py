from flask import Flask, request, render_template, jsonify, send_file, session, redirect, url_for, flash
import os
import base64
import requests
from werkzeug.utils import secure_filename
import json
from PIL import Image
import io
from datetime import datetime, timedelta
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
import re
import psycopg2
from functools import wraps
import config  # 导入配置文件
import bcrypt

app = Flask(__name__)
app.secret_key = config.SECRET_KEY

# Database Configuration
DB_URI = config.DB_URI

# 配置
app.config['MAX_CONTENT_LENGTH'] = config.MAX_CONTENT_LENGTH
app.config['UPLOAD_FOLDER'] = config.UPLOAD_FOLDER

# 设置日志
log_level = config.LOG_LEVEL
logging.basicConfig(level=getattr(logging, log_level))
logger = logging.getLogger(__name__)

# Usage Log File
USAGE_LOG_FILE = config.USAGE_LOG_FILE

# 创建上传目录
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# 转换次数统计文件
STATS_FILE = config.STATS_FILE

# 支持的图片格式
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'}

# --- Database & Auth Functions ---

def get_db_connection():
    try:
        conn = psycopg2.connect(DB_URI)
        return conn
    except Exception as e:
        logger.error(f"Database connection error: {e}")
        return None

def check_login(email, password):
    conn = get_db_connection()
    if not conn:
        return False, "数据库连接失败"
    
    try:
        cur = conn.cursor()
        cur.execute("SELECT password, active FROM auth WHERE email = %s", (email,))
        user = cur.fetchone()
        
        if not user:
            return False, "用户不存在"
        
        db_password, active = user
        
        # 验证密码 (支持 bcrypt 哈希和明文)
        password_valid = False
        try:
            # 检查是否为 bcrypt 哈希 (通常以 $2b$, $2a$, $2y$ 开头)
            if db_password.startswith(('$2b$', '$2a$', '$2y$')):
                if bcrypt.checkpw(password.encode('utf-8'), db_password.encode('utf-8')):
                    password_valid = True
            else:
                # 明文匹配 (用于测试账户或旧数据)
                if db_password == password:
                    password_valid = True
        except Exception as e:
            logger.error(f"Password verification error: {e}")
            # 出错时尝试明文匹配作为后备
            if db_password == password:
                password_valid = True

        if not password_valid:
            return False, "密码错误"
        
        if not active:
            return False, "账户未激活"
            
        return True, "登录成功"
    except Exception as e:
        logger.error(f"Login error: {e}")
        return False, f"登录处理错误: {str(e)}"
    finally:
        if conn:
            conn.close()

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_email' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# --- Logging Functions ---

def log_user_action(email, success):
    try:
        timestamp = datetime.now().isoformat()
        status = "Success" if success else "Failed"
        log_entry = f"{timestamp}|{email}|{status}\n"
        
        with open(USAGE_LOG_FILE, 'a', encoding='utf-8') as f:
            f.write(log_entry)
    except Exception as e:
        logger.error(f"Failed to write usage log: {e}")

def mask_email(email):
    if not email or '@' not in email:
        return email
    try:
        username, domain = email.split('@')
        if len(username) <= 2:
            masked_username = username[0] + "***"
        else:
            masked_username = username[:2] + "***" + username[-1]
        return f"{masked_username}@{domain}"
    except:
        return email

def get_display_logs():
    if not os.path.exists(USAGE_LOG_FILE):
        return []
    
    logs = []
    try:
        with open(USAGE_LOG_FILE, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # Parse lines
        parsed_logs = []
        for line in lines:
            try:
                parts = line.strip().split('|')
                if len(parts) >= 3:
                    timestamp_str = parts[0]
                    email = parts[1]
                    status = parts[2]
                    dt = datetime.fromisoformat(timestamp_str)
                    parsed_logs.append({
                        'dt': dt,
                        'email': email,
                        'status': status,
                        'original_line': line
                    })
            except Exception as e:
                continue
        
        # Sort by time descending
        parsed_logs.sort(key=lambda x: x['dt'], reverse=True)
        
        # Merge logs
        merged_logs = []
        if not parsed_logs:
            return []
            
        current_group = [parsed_logs[0]]
        
        for i in range(1, len(parsed_logs)):
            log = parsed_logs[i]
            prev = current_group[-1]
            
            # Check if same user and status within 30 minutes
            time_diff = abs((prev['dt'] - log['dt']).total_seconds())
            if log['email'] == prev['email'] and log['status'] == prev['status'] and time_diff <= 1800:
                current_group.append(log)
            else:
                # Process current group -> take the latest one (which is the first in group due to sort)
                latest_log = current_group[0]
                count = len(current_group)
                display_time = latest_log['dt'].strftime("%Y-%m-%d %H:%M:%S")
                masked = mask_email(latest_log['email'])
                
                msg = f"用户 {masked} 在 {display_time} 使用转换功能 {'成功' if latest_log['status'] == 'Success' else '失败'}"
                if count > 1:
                    msg += f" (含近30分钟内 {count} 次记录)"
                
                merged_logs.append(msg)
                current_group = [log]
        
        # Process last group
        if current_group:
            latest_log = current_group[0]
            count = len(current_group)
            display_time = latest_log['dt'].strftime("%Y-%m-%d %H:%M:%S")
            masked = mask_email(latest_log['email'])
            
            msg = f"用户 {masked} 在 {display_time} 使用转换功能 {'成功' if latest_log['status'] == 'Success' else '失败'}"
            if count > 1:
                msg += f" (含近30分钟内 {count} 次记录)"
            
            merged_logs.append(msg)
            
        return merged_logs
        
    except Exception as e:
        logger.error(f"Error reading logs: {e}")
        return []

# --- Existing Helper Functions ---

def get_conversion_count():
    """获取转换次数"""
    try:
        if os.path.exists(STATS_FILE):
            with open(STATS_FILE, 'r') as f:
                return int(f.read().strip())
        return 0
    except:
        return 0

def increment_conversion_count():
    """增加转换次数"""
    try:
        count = get_conversion_count() + 1
        with open(STATS_FILE, 'w') as f:
            f.write(str(count))
        return count
    except:
        return get_conversion_count()

def get_api_config():
    """获取API配置"""
    return {
        'api_key': config.OPENAI_API_KEY,
        'api_base': config.OPENAI_API_BASE,
        'model': config.OPENAI_MODEL,
        'max_tokens': config.MODEL_MAX_TOKENS,
        'temperature': config.MODEL_TEMPERATURE,
        'image_max_size': config.IMAGE_MAX_SIZE,
        'image_quality': config.IMAGE_QUALITY
    }

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def optimize_image(image_data, config):
    """优化图片大小和质量"""
    try:
        # 解码base64图片
        if ',' in image_data:
            image_data = image_data.split(',')[1]
        
        img_bytes = base64.b64decode(image_data)
        img = Image.open(io.BytesIO(img_bytes))
        
        # 转换为RGB模式（如果需要）
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # 调整图片大小
        max_size = (config['image_max_size'], config['image_max_size'])
        img.thumbnail(max_size, Image.Resampling.LANCZOS)
        
        # 保存为优化后的JPEG
        output = io.BytesIO()
        img.save(output, format='JPEG', quality=config['image_quality'], optimize=True)
        
        # 转换回base64
        optimized_b64 = base64.b64encode(output.getvalue()).decode('utf-8')
        
        return optimized_b64
    except Exception as e:
        logger.error(f"图片优化失败: {e}")
        return image_data.split(',')[1] if ',' in image_data else image_data

def get_prompt_text():
    """获取AI提示词"""
    return """请识别这张图片中的数学公式，并将其转换为LaTeX代码。要求：
1. 只返回LaTeX代码，不要其他解释
2. 如果是行内公式，用 $...$ 包围
3. 如果是行间公式，用 $$...$$ 包围
4. 确保LaTeX语法正确，可以在Word中直接使用
5. 如果图片中有多个公式，请分别转换，每个公式单独一行
6. 特别注意分数、根号、积分、求和等复杂符号的正确表示
7. 如果无法识别出数学公式，请返回"无法识别数学公式"

示例格式：
简单公式：$x^2 + y^2 = z^2$
复杂公式：$$\\int_{0}^{\\infty} e^{-x^2} dx = \\frac{\\sqrt{\\pi}}{2}$$
分数公式：$$\\frac{a + b}{c - d} = \\frac{\\sqrt{x}}{y^2}$$
求和公式：$$\\sum_{i=1}^{n} x_i = \\frac{n(n+1)}{2}$$"""

def call_ai_api(image_base64, user_email=None):
    """调用AI API进行数学公式识别"""
    config = get_api_config()
    
    if not config['api_key']:
        if user_email: log_user_action(user_email, False)
        return {"success": False, "error": "API密钥未配置"}
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config['api_key']}"
    }
    
    # 优化图片
    optimized_image = optimize_image(image_data=image_base64, config=config)
    
    # 构建请求payload
    payload = {
        "model": config['model'],
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": get_prompt_text()
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{optimized_image}",
                            "detail": "high"
                        }
                    }
                ]
            }
        ],
        "max_tokens": config['max_tokens'],
        "temperature": config['temperature']
    }
    
    try:
        logger.info(f"正在调用API: {config['api_base']}, 模型: {config['model']}")
        
        response = requests.post(
            f"{config['api_base']}/chat/completions",
            headers=headers, 
            json=payload, 
            timeout=60
        )
        
        logger.info(f"API请求状态码: {response.status_code}")
        
        if response.status_code != 200:
            error_text = response.text[:500] if response.text else "无响应内容"
            logger.error(f"API请求失败，状态码: {response.status_code}, 响应: {error_text}")
            if user_email: log_user_action(user_email, False)
            return {"success": False, "error": f"API请求失败，状态码: {response.status_code}"}
        
        result = response.json()
        
        if 'choices' not in result or len(result['choices']) == 0:
            logger.error(f"API响应格式错误: {result}")
            if user_email: log_user_action(user_email, False)
            return {"success": False, "error": "API响应格式错误，没有找到choices字段"}
        
        latex_code = result['choices'][0]['message']['content'].strip()
        
        # 清理和格式化LaTeX代码
        latex_code = clean_latex_output(latex_code)
        
        # 转换成功，增加计数
        count = increment_conversion_count()
        
        logger.info(f"API调用成功，返回LaTeX代码长度: {len(latex_code)}，总转换次数: {count}")
        if user_email: log_user_action(user_email, True)
        return {"success": True, "latex": latex_code, "total_conversions": count}
    
    except Exception as e:
        logger.error(f"处理失败: {e}")
        if user_email: log_user_action(user_email, False)
        return {"success": False, "error": f"处理失败: {str(e)}"}

def clean_latex_output(latex_code):
    """清理和格式化LaTeX输出"""
    # 移除可能的markdown代码块标记
    latex_code = latex_code.replace('```latex', '').replace('```', '')
    latex_code = latex_code.replace('```LaTeX', '').replace('```tex', '')
    
    # 移除多余的空白字符
    latex_code = latex_code.strip()
    
    # 如果返回了"无法识别"的消息，直接返回
    if "无法识别" in latex_code or "cannot" in latex_code.lower() or "unable" in latex_code.lower():
        return latex_code
    
    # 确保公式有适当的包围符号
    lines = latex_code.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        if line and not line.startswith('$') and not line.startswith('\\['):
            # 检查是否包含LaTeX数学符号
            math_symbols = ['\\frac', '\\int', '\\sum', '\\sqrt', '^', '_', 
                          '\\alpha', '\\beta', '\\gamma', '\\delta', '\\theta',
                          '\\pi', '\\sigma', '\\lambda', '\\mu', '\\nu',
                          '\\partial', '\\nabla', '\\infty', '\\lim']
            
            if any(symbol in line for symbol in math_symbols):
                # 复杂公式用显示模式
                if not line.startswith('$$'):
                    line = f'$${line}$$'
            else:
                # 简单公式用行内模式
                if not line.startswith('$'):
                    line = f'${line}$'
        cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def create_word_document_with_formula(latex_code):
    """创建包含数学公式的Word文档"""
    try:
        # 创建新的Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('数学公式转换结果', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加生成时间
        time_para = doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        time_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加分隔线
        separator_para = doc.add_paragraph('=' * 50)
        separator_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 处理多行LaTeX代码
        formulas = [line.strip() for line in latex_code.split('\n') if line.strip()]
        
        for i, formula in enumerate(formulas, 1):
            if formula and not formula.startswith('#'):
                # 添加公式标题
                formula_title = doc.add_heading(f'公式 {i}', level=2)
                formula_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # 添加LaTeX代码显示
                latex_heading = doc.add_paragraph()
                latex_run = latex_heading.add_run('LaTeX代码: ')
                latex_run.bold = True
                latex_run.font.size = Pt(12)
                
                latex_code_para = doc.add_paragraph()
                latex_code_run = latex_code_para.add_run(formula)
                latex_code_run.font.name = 'Consolas'
                latex_code_run.font.size = Pt(11)
                # 设置背景色（浅灰色）
                latex_code_para.paragraph_format.space_before = Pt(6)
                latex_code_para.paragraph_format.space_after = Pt(6)
                
                # 添加可编辑的公式区域
                equation_heading = doc.add_paragraph()
                equation_run = equation_heading.add_run('可编辑公式: ')
                equation_run.bold = True
                equation_run.font.size = Pt(12)
                
                # 创建空白段落用于插入公式
                equation_para = doc.add_paragraph()
                equation_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # 清理公式
                clean_formula = formula.strip()
                if clean_formula.startswith('$$') and clean_formula.endswith('$$'):
                    clean_formula = clean_formula[2:-2]
                elif clean_formula.startswith('$') and clean_formula.endswith('$'):
                    clean_formula = clean_formula[1:-1]
                
                # 添加占位符文本，用户需要手动替换为公式
                placeholder_run = equation_para.add_run(f"[请选中此文本，按Alt+=，然后输入: {clean_formula}]")
                placeholder_run.font.name = 'Cambria Math'
                placeholder_run.font.size = Pt(14)
                placeholder_run.italic = True
                
                # 添加说明文字
                instruction_para = doc.add_paragraph()
                instruction_text = (
                    f"📝 使用步骤：\n"
                    f"1. 选中上方的占位符文本\n"
                    f"2. 按 Alt + = 键打开公式编辑器\n"
                    f"3. 粘贴LaTeX代码: {clean_formula}\n"
                    f"4. 按回车键完成公式插入"
                )
                instruction_run = instruction_para.add_run(instruction_text)
                instruction_run.font.size = Pt(10)
                instruction_run.font.color.rgb = docx.shared.RGBColor(108, 117, 125)  # 灰色文字
                
                # 添加分隔空行
                doc.add_paragraph()
        
        # 添加页面分隔
        doc.add_page_break()
        
        # 添加详细使用说明
        doc.add_heading('📖 详细使用说明', level=1)
        
        # 基本步骤
        basic_steps_heading = doc.add_heading('基本操作步骤', level=2)
        basic_steps = [
            "1. 找到文档中的占位符文本（灰色斜体部分）",
            "2. 用鼠标选中整个占位符文本",
            "3. 按键盘上的 Alt + = 组合键",
            "4. Word会自动打开公式编辑器",
            "5. 在公式编辑器中，点击顶部的 'LaTeX' 按钮",
            "6. 粘贴对应的LaTeX代码",
            "7. 按回车键，公式会自动转换并显示"
        ]
        
        for step in basic_steps:
            step_para = doc.add_paragraph(step)
            step_para.paragraph_format.left_indent = Inches(0.25)
        
        # 注意事项
        doc.add_heading('⚠️ 注意事项', level=2)
        notes = [
            "• 确保使用Microsoft Word 2016或更新版本",
            "• 如果看不到LaTeX选项，请更新Office软件",
            "• 复杂公式可能需要分步骤输入",
            "• 建议先在简单公式上练习操作步骤",
            "• 如果公式显示异常，可以撤销后重新操作"
        ]
        
        for note in notes:
            note_para = doc.add_paragraph(note)
            note_para.paragraph_format.left_indent = Inches(0.25)
        
        # 快捷键说明
        doc.add_heading('⌨️ 常用快捷键', level=2)
        shortcuts = [
            "Alt + = : 插入公式",
            "Ctrl + Z : 撤销操作", 
            "Ctrl + Y : 重做操作",
            "Esc : 退出公式编辑模式"
        ]
        
        for shortcut in shortcuts:
            shortcut_para = doc.add_paragraph(shortcut)
            shortcut_para.paragraph_format.left_indent = Inches(0.25)
        
        # 添加页脚
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "数学公式图片转LaTeX工具 | Powered by AI | https://github.com/your-project"
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        return doc
        
    except Exception as e:
        logger.error(f"创建Word文档失败: {e}")
        raise e

# --- Routes ---

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        
        success, message = check_login(email, password)
        if success:
            session['user_email'] = email
            return redirect(url_for('index'))
        else:
            flash(message)
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('user_email', None)
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    count = get_conversion_count()
    logs = get_display_logs()
    return render_template('index.html', total_conversions=count, usage_logs=logs, user_email=session.get('user_email'))

@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': '没有上传文件'})
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'success': False, 'error': '没有选择文件'})
    
    if file and allowed_file(file.filename):
        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # 转换为base64
            with open(filepath, "rb") as image_file:
                image_base64 = base64.b64encode(image_file.read()).decode('utf-8')
            
            # 调用API
            result = call_ai_api(image_base64, user_email=session.get('user_email'))
            
            # 删除临时文件
            os.remove(filepath)
            
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"文件处理错误: {e}")
            if 'user_email' in session:
                log_user_action(session['user_email'], False)
            return jsonify({'success': False, 'error': f'处理文件时出错: {str(e)}'})
    
    return jsonify({'success': False, 'error': '不支持的文件格式'})

@app.route('/upload_base64', methods=['POST'])
@login_required
def upload_base64():
    """处理粘贴的图片数据"""
    try:
        data = request.json
        image_data = data.get('image')
        
        if not image_data:
            return jsonify({'success': False, 'error': '没有图片数据'})
        
        logger.info("接收到图片数据，开始转换")
        
        # 调用API
        result = call_ai_api(image_data, user_email=session.get('user_email'))
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Base64处理错误: {e}")
        if 'user_email' in session:
            log_user_action(session['user_email'], False)
        return jsonify({'success': False, 'error': f'处理数据时出错: {str(e)}'})

@app.route('/download_word', methods=['POST'])
@login_required
def download_word():
    """生成并下载包含公式的Word文档"""
    try:
        data = request.json
        latex_code = data.get('latex', '')
        
        if not latex_code:
            return jsonify({'success': False, 'error': '没有LaTeX代码'})
        
        logger.info("开始生成Word文档")
        
        # 创建Word文档
        doc = create_word_document_with_formula(latex_code)
        
        # 使用临时文件，不保存到服务器
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            tmp_file_path = tmp_file.name
        
        # 生成文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"math_formula_{timestamp}.docx"
        
        logger.info(f"Word文档生成成功: {filename}")
        
        response = send_file(
            tmp_file_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # 在响应后清理临时文件
        @response.call_on_close
        def remove_file():
            try:
                os.unlink(tmp_file_path)
            except Exception as e:
                logger.error(f"删除临时文件失败: {e}")
        
        return response
        
    except Exception as e:
        logger.error(f"生成Word文档失败: {e}")
        return jsonify({'success': False, 'error': f'生成Word文档失败: {str(e)}'})

@app.route('/stats')
def get_stats():
    """获取统计信息"""
    return jsonify({
        'total_conversions': get_conversion_count(),
        'timestamp': datetime.now().isoformat()
    })

@app.route('/health')
def health_check():
    """健康检查接口"""
    try:
        config = get_api_config()
        return jsonify({
            'status': 'healthy', 
            'timestamp': datetime.now().isoformat(),
            'version': '1.0.0',
            'api_configured': bool(config['api_key'])
        })
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'error': str(e)
        }), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
